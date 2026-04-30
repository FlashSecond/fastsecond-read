"""
Word文档读取器 V2 - 基于版面特征的智能分章
使用段落样式、字体大小、对齐方式等特征进行章节识别

智能分章规则：
1. 章节标题：Heading 1-2 样式，或字号显著大于正文
2. 章节内小标题：Heading 3-9 样式，或数字列表格式
3. 正文：Normal 样式，或字号接近正文基准
4. 内容归属：章节标题后的所有内容归属于该章节
"""
from .base import FileReader
from pathlib import Path
from typing import List, Dict
from core.document import Document, Chapter, ContentBlock, ContentType, TextStyle


class DocxReader(FileReader):
    """
    Word文档读取器 V2 - 基于版面特征的智能分章
    
    基于样式特征的智能分章：
    - 章节标题：Heading 1-2，字号显著大于正文
    - 章节内小标题：Heading 3+，或数字列表格式
    - 正文：Normal 样式，字号接近正文基准
    """
    
    def supports(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in ['.docx', '.doc']
    
    def read(self, file_path: str) -> Document:
        """读取Word文档并返回结构化文档"""
        file_info = self.get_file_info(file_path)
        
        try:
            from docx import Document as DocxDocument
            
            doc = Document(
                file_path=file_info["path"],
                file_name=file_info["name"],
                file_format="docx"
            )
            
            docx_doc = DocxDocument(file_path)
            
            # 提取元数据
            doc.title = docx_doc.core_properties.title
            doc.author = docx_doc.core_properties.author
            
            # 提取所有段落块
            all_blocks = []
            for para in docx_doc.paragraphs:
                block = self._create_block_from_paragraph(para)
                if block:
                    all_blocks.append(block)
            
            # 分析字体统计
            font_stats = self._analyze_font_stats(all_blocks)
            
            # 分类文本块
            self._classify_blocks(all_blocks, font_stats)
            
            # 检测章节
            doc.chapters = self._detect_chapters(all_blocks)
            doc.total_chapters = len(doc.chapters)
            doc.total_words = sum(len(b["text"]) for b in all_blocks)
            
            return doc
            
        except ImportError:
            print("python-docx not installed. Please install: pip install python-docx")
            return self._create_empty_doc(file_info)
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {e}")
            return self._create_empty_doc(file_info)
    
    def _create_block_from_paragraph(self, para) -> Dict:
        """从段落创建内容块字典"""
        text = para.text.strip()
        if not text:
            return None
        
        # 提取样式信息
        style_info = self._extract_style_info(para)
        
        # 检测对齐方式
        alignment = self._detect_alignment(para)
        
        # 检测样式级别
        style_level = self._detect_style_level(para)
        
        return {
            "text": text,
            "style_name": para.style.name if para.style else "Normal",
            "style_level": style_level,
            "font_size": style_info.get("font_size"),
            "is_bold": style_info.get("is_bold", False),
            "is_italic": style_info.get("is_italic", False),
            "alignment": alignment,
            "is_centered": alignment == "center",
            "is_left_aligned": alignment == "left",
            "content_block": None  # 稍后创建
        }
    
    def _extract_style_info(self, para) -> Dict:
        """提取段落样式信息"""
        info = {
            "font_size": None,
            "is_bold": False,
            "is_italic": False,
            "font_name": None
        }
        
        if not para.runs:
            return info
        
        font_sizes = []
        for run in para.runs:
            if run.bold:
                info["is_bold"] = True
            if run.italic:
                info["is_italic"] = True
            if run.font.size:
                # 转换为磅
                size_pt = run.font.size.pt if hasattr(run.font.size, 'pt') else None
                if size_pt:
                    font_sizes.append(size_pt)
            if run.font.name:
                info["font_name"] = run.font.name
        
        # 使用平均字体大小
        if font_sizes:
            info["font_size"] = sum(font_sizes) / len(font_sizes)
        
        return info
    
    def _detect_style_level(self, para) -> int:
        """检测段落样式级别"""
        import re
        
        style_name = para.style.name.lower() if para.style else "normal"
        
        # Heading 样式
        if 'heading' in style_name or '标题' in style_name:
            match = re.search(r'(\d+)', style_name)
            if match:
                return int(match.group(1))
            return 1
        
        return 0
    
    def _detect_alignment(self, para) -> str:
        """检测段落对齐方式"""
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: 'left',
                WD_ALIGN_PARAGRAPH.CENTER: 'center',
                WD_ALIGN_PARAGRAPH.RIGHT: 'right',
                WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
            }
            
            return alignment_map.get(para.alignment, 'left')
        except:
            return 'left'
    
    def _analyze_font_stats(self, blocks: List[Dict]) -> Dict:
        """分析全局字体统计信息"""
        if not blocks:
            return {}
        
        font_sizes = [b["font_size"] for b in blocks if b.get("font_size")]
        
        if not font_sizes:
            return {"median_size": 12, "q3_size": 14}
        
        font_sizes.sort()
        n = len(font_sizes)
        
        return {
            "min_size": font_sizes[0],
            "max_size": font_sizes[-1],
            "median_size": font_sizes[n // 2],
            "mean_size": sum(font_sizes) / n,
            "q1_size": font_sizes[n // 4],
            "q3_size": font_sizes[3 * n // 4],
        }
    
    def _classify_blocks(self, blocks: List[Dict], font_stats: Dict):
        """
        根据版面特征分类文本块
        
        分类体系：
        - 章节标题: Heading 1-2 或字号显著大于正文
        - 小标题: Heading 3+ 或数字列表格式
        - 正文: Normal 样式
        """
        import re
        
        median_size = font_stats.get("median_size", 12)
        q3_size = font_stats.get("q3_size", 14)
        
        for block in blocks:
            text = block["text"]
            style_level = block.get("style_level", 0)
            font_size = block.get("font_size") or median_size
            is_bold = block.get("is_bold", False)
            is_centered = block.get("is_centered", False)
            text_len = len(text)
            
            # 初始化
            is_chapter_title = False
            is_subheading = False
            subheading_level = 0
            
            # ========== 章节标题判定 ==========
            # 条件1: Heading 1-2 样式
            if style_level in [1, 2]:
                is_chapter_title = True
            # 条件2: 字号显著大于正文 + 居中/粗体 + 短文本
            elif font_size > q3_size * 1.15 and text_len < 50:
                if is_centered or is_bold:
                    is_chapter_title = True
            # 条件3: 匹配章节标题格式
            elif self._is_chapter_title_pattern(text) and text_len < 60:
                if style_level > 0 or font_size > median_size * 1.1:
                    is_chapter_title = True
            
            # ========== 小标题判定（章节内） ==========
            if not is_chapter_title:
                # 条件1: Heading 3+ 样式
                if style_level >= 3:
                    is_subheading = True
                    subheading_level = style_level
                # 条件2: 数字列表格式 + 字号略大
                elif self._is_numbered_subheading(text):
                    if median_size * 1.05 <= font_size <= median_size * 1.3:
                        is_subheading = True
                        subheading_level = 2
                # 条件3: 粗体 + 短文本
                elif is_bold and text_len < 40:
                    is_subheading = True
                    subheading_level = 3
            
            block["is_chapter_title"] = is_chapter_title
            block["is_subheading"] = is_subheading
            block["is_body"] = not is_chapter_title and not is_subheading
            block["subheading_level"] = subheading_level
    
    def _is_chapter_title_pattern(self, text: str) -> bool:
        """判断是否为章节标题格式"""
        import re
        
        patterns = [
            r'^[第][\d一二三四五六七八九十百千万]+[章节篇回]',
            r'^Chapter\s+\d+',
            r'^Part\s+\d+',
            r'^Section\s+\d+',
        ]
        
        for pattern in patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        return False
    
    def _is_numbered_subheading(self, text: str) -> bool:
        """判断是否为数字编号的小标题"""
        import re
        
        patterns = [
            r'^\d{1,3}[\.\、\s]\s*',
            r'^\d{2,3}\s+',
            r'^[（(]\d{1,3}[)）]\s*',
            r'^[\u4e00-\u9fa5]{1,3}[、\.\s]',
            r'^[（(][\u4e00-\u9fa5]{1,3}[)）]',
        ]
        
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        return False
    
    def _detect_chapters(self, blocks: List[Dict]) -> List[Chapter]:
        """
        根据章节标题检测章节边界
        
        规则：
        - 只有 is_chapter_title 开始新章节
        - 小标题和正文都归属于当前章节
        """
        chapters = []
        current_chapter_blocks = []
        chapter_index = 0
        current_title = None
        
        for block in blocks:
            if block.get("is_chapter_title"):
                # 保存上一个章节
                if current_chapter_blocks and current_title:
                    chapter = self._create_chapter(
                        chapter_index, 
                        current_title, 
                        current_chapter_blocks
                    )
                    if chapter:
                        chapters.append(chapter)
                        chapter_index += 1
                
                # 开始新章节
                current_title = block["text"]
                current_chapter_blocks = [block]
            else:
                # 归属于当前章节
                current_chapter_blocks.append(block)
        
        # 保存最后一个章节
        if current_chapter_blocks and current_title:
            chapter = self._create_chapter(
                chapter_index,
                current_title,
                current_chapter_blocks
            )
            if chapter:
                chapters.append(chapter)
        
        # 如果没有检测到章节，将所有内容作为一个章节
        if not chapters and blocks:
            title = blocks[0]["text"][:50] if blocks else "正文"
            chapter = self._create_chapter(0, title, blocks)
            if chapter:
                chapters.append(chapter)
        
        return chapters
    
    def _create_chapter(self, index: int, title: str, blocks: List[Dict]) -> Chapter:
        """从文本块创建章节"""
        if not blocks:
            return None
        
        content_blocks = []
        
        # 第一个块是章节标题，跳过
        body_blocks = blocks[1:] if len(blocks) > 1 else blocks
        
        for block in body_blocks:
            text = block.get("text", "")
            if not text:
                continue
            
            # 确定内容类型
            if block.get("is_subheading"):
                content_type = ContentType.HEADING
                level = block.get("subheading_level", 2)
            elif block.get("is_body"):
                content_type = ContentType.PARAGRAPH
                level = 0
            else:
                content_type = ContentType.PARAGRAPH
                level = 0
            
            # 创建样式
            style = TextStyle(
                bold=block.get("is_bold", False),
                italic=block.get("is_italic", False),
                font_size=block.get("font_size"),
                alignment=block.get("alignment")
            )
            
            content_block = ContentBlock(
                type=content_type,
                text=text,
                level=level,
                style=style
            )
            content_blocks.append(content_block)
        
        # 计算统计信息
        word_count = sum(len(b.text) for b in content_blocks)
        paragraph_count = sum(
            1 for b in content_blocks 
            if b.type == ContentType.PARAGRAPH
        )
        
        return Chapter(
            index=index + 1,
            title=title[:100],
            level=1,
            content_blocks=content_blocks,
            word_count=word_count,
            paragraph_count=paragraph_count
        )
    
    def _create_empty_doc(self, file_info: dict) -> Document:
        """创建空文档"""
        return Document(
            file_path=file_info["path"],
            file_name=file_info["name"],
            file_format="docx"
        )
